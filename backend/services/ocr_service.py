"""
OCR Service — Sprint 8 Pluggable Adapter
=========================================
Liztek Procure-AI: Replaces the regex _nlp_extract() in po_intake_agent.py
and _ocr_extract_invoice() in invoice_capture_agent.py.

Architecture
------------
IOCRService (ABC)
  ├── RegexOCRService        — regex-based extraction (default, no API key)
  ├── MindeeOCRService       — Mindee API (MINDEE_API_KEY required)
  └── AWSTextractService     — AWS Textract via boto3 (AWS credentials required)

Factory
-------
get_ocr_service() reads OCR_PROVIDER env var:
  'regex'    → RegexOCRService (default)
  'mindee'   → MindeeOCRService
  'textract' → AWSTextractService

Environment Variables
---------------------
OCR_PROVIDER=regex          (default — no API key needed)
OCR_PROVIDER=mindee         (requires MINDEE_API_KEY)
OCR_PROVIDER=textract       (requires AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY, AWS_REGION)
MINDEE_API_KEY=             (Mindee API key)
AWS_ACCESS_KEY_ID=          (AWS access key)
AWS_SECRET_ACCESS_KEY=      (AWS secret key)
AWS_REGION=us-east-1        (AWS region, default: us-east-1)
"""

from __future__ import annotations

import logging
import os
import re
from abc import ABC, abstractmethod
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)


# ── Abstract Base ─────────────────────────────────────────────────────────────

class IOCRService(ABC):
    """
    Abstract base for all OCR / document-extraction providers.

    Every implementation must support extracting both PO and Invoice fields
    from raw document text (and optionally from a file path for binary uploads).
    """

    @abstractmethod
    def extract_po(self, raw_text: str, file_path: Optional[str] = None) -> dict:
        """
        Extract PO fields from raw document text or file.

        Returns dict with keys (all optional):
          po_number, vendor, po_date, total_amount, currency,
          line_items, estimated_line_count
        """

    @abstractmethod
    def extract_invoice(self, raw_text: str, file_path: Optional[str] = None) -> dict:
        """
        Extract invoice fields from raw document text or file.

        Returns dict with keys (all optional):
          invoice_number, po_reference, po_number, vendor, amount,
          total_amount, currency, invoice_date, due_date, tax_amount,
          payment_terms, line_items
        """


# ── Implementation 1: RegexOCRService (default / fallback) ───────────────────

class RegexOCRService(IOCRService):
    """
    Regex-based field extraction — the original logic from po_intake_agent
    and invoice_capture_agent, promoted to a first-class service.

    No external dependencies. Works without any API key.
    Used as the fallback when other providers fail.
    """

    # ── PO extraction ─────────────────────────────────────────────────────────

    def extract_po(self, raw_text: str, file_path: Optional[str] = None) -> dict:
        """
        Extract PO fields using regex patterns.

        Patterns cover common PO document layouts:
          - PO number: PO-1234, PO/2025/001, 4500012345, "Purchase Order: xxxx"
          - Vendor / supplier name
          - Date in common formats
          - Total amount with optional currency prefix
          - Currency code
          - Line item count estimate
        """
        text = raw_text or ''
        fields: Dict[str, Any] = {}

        # PO number: PO-1234, PO/2025/001, PO-ALPHA-001, 4500012345, "purchase order: XXX"
        # Try explicit label first (most reliable), then bare PO-prefixed pattern
        po_label = re.search(
            r'(?:p\.?o\.?\s*(?:no|number|#|num)?|purchase\s*order)[:\s#]+([A-Za-z0-9][\w\-/]{2,30})',
            text, re.IGNORECASE
        )
        po_bare = re.search(r'\bPO[-/][\w\-]{2,30}\b', text, re.IGNORECASE)
        if po_label:
            fields['po_number'] = po_label.group(1).strip()
        elif po_bare:
            fields['po_number'] = po_bare.group(0).strip()

        # Vendor / supplier — stop capture at common keyword boundaries
        vendor_match = re.search(
            r'(?:vendor|supplier|from)[:\s]+([A-Za-z][^,\n]{2,60}?)(?=\s*(?:date|total|amount|currency|po|order|$|\n))',
            text, re.IGNORECASE
        )
        if vendor_match:
            fields['vendor'] = vendor_match.group(1).strip()

        # Date (various formats)
        date_match = re.search(
            r'\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}-\d{2}-\d{2})\b', text
        )
        if date_match:
            fields['po_date'] = date_match.group(1)

        # Total amount
        amount_match = re.search(
            r'(?:total|amount|grand total)[:\s]*(?:USD|EUR|GBP|AED|SAR|INR)?\s*([\d,]+\.?\d*)',
            text, re.IGNORECASE
        )
        if amount_match:
            fields['total_amount'] = amount_match.group(1).replace(',', '')

        # Currency
        curr_match = re.search(r'\b(USD|EUR|GBP|AED|SAR|INR|JPY|CNY)\b', text)
        if curr_match:
            fields['currency'] = curr_match.group(1)

        # Line item count estimate
        line_count = len(re.findall(r'^\s*\d+[\s.)\-]+\S', text, re.MULTILINE))
        if line_count > 0:
            fields['estimated_line_count'] = line_count

        logger.debug("[RegexOCRService] extract_po — fields found: %s", list(fields.keys()))
        return fields

    # ── Invoice extraction ────────────────────────────────────────────────────

    def extract_invoice(self, raw_text: str, file_path: Optional[str] = None) -> dict:
        """
        Extract invoice fields using regex patterns.

        Patterns cover common invoice layouts:
          - Invoice number, PO reference, vendor name
          - Invoice date, due date
          - Total amount, currency, tax amount
          - Payment terms (Net 30, etc.)
        """
        text = raw_text or ''
        fields: Dict[str, Any] = {}

        # Invoice number — match "Invoice No: INV-001", "Invoice INV-001", "TAX INVOICE INV-001"
        # Also matches "INVOICE BILL/2025/00001" (Odoo-style)
        inv_match = re.search(
            r'(?:(?:tax\s+)?invoice\s*(?:no|number|#|num)?[:\s#]*)((?:INV|BILL|INV)[-/][\w\-/]+|[\w\-/]*\d{3,}[\w\-/]*)',
            text, re.IGNORECASE
        )
        if inv_match:
            val = inv_match.group(1).strip()
            if val and val.lower() not in ('no', 'number', 'num', '#'):
                fields['invoice_number'] = val

        # PO reference — "PO Reference: PO-001", "PO No: PO-001", "PO: PO-001"
        # Put 'reference' before 'ref' to avoid partial match on 'reference'
        po_match = re.search(
            r'(?:po\s*(?:no|number|#|reference|ref)\s*[:\s#]*|purchase\s*order\s*(?:no|number|#|ref|reference)[:\s#]*)([A-Za-z0-9][\w\-/]{2,30})',
            text, re.IGNORECASE
        )
        if po_match:
            fields['po_number'] = po_match.group(1).strip()
            fields['po_reference'] = fields['po_number']

        # Vendor
        vendor_match = re.search(
            r'(?:vendor|supplier|billed?\s*by|from)[:\s]+([A-Za-z][^\n,]{2,50})',
            text, re.IGNORECASE
        )
        if vendor_match:
            fields['vendor'] = vendor_match.group(1).strip()

        # Invoice date
        date_match = re.search(
            r'(?:invoice\s*date|date)[:\s]+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}-\d{2}-\d{2})',
            text, re.IGNORECASE
        )
        if date_match:
            fields['invoice_date'] = date_match.group(1)

        # Due date
        due_match = re.search(
            r'(?:due\s*date|payment\s*due)[:\s]+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}-\d{2}-\d{2})',
            text, re.IGNORECASE
        )
        if due_match:
            fields['due_date'] = due_match.group(1)

        # Total amount
        total_match = re.search(
            r'(?:total\s*(?:amount|due|payable)|amount\s*due)[:\s]*(?:USD|EUR|GBP|AED|SAR|INR)?\s*([\d,]+\.?\d*)',
            text, re.IGNORECASE
        )
        if total_match:
            fields['total_amount'] = total_match.group(1).replace(',', '')
            fields['amount'] = fields['total_amount']

        # Currency
        curr_match = re.search(r'\b(USD|EUR|GBP|AED|SAR|INR|JPY|CNY)\b', text)
        if curr_match:
            fields['currency'] = curr_match.group(1)

        # Tax amount
        tax_match = re.search(
            r'(?:tax|vat|gst)[:\s]*(?:USD|EUR|GBP|AED|SAR|INR)?\s*([\d,]+\.?\d*)',
            text, re.IGNORECASE
        )
        if tax_match:
            fields['tax_amount'] = tax_match.group(1).replace(',', '')

        # Payment terms (Net 30, Net 60, etc.)
        terms_match = re.search(r'\b(Net\s*\d+|Immediate|Due on receipt)\b', text, re.IGNORECASE)
        if terms_match:
            fields['payment_terms'] = terms_match.group(1)

        logger.debug("[RegexOCRService] extract_invoice — fields found: %s", list(fields.keys()))
        return fields


# ── Implementation 2: MindeeOCRService ───────────────────────────────────────

class MindeeOCRService(IOCRService):
    """
    Mindee API-based document extraction.

    Uses the Mindee Purchase Receipts API to extract structured fields.
    Maps Mindee response fields to the Procure-AI neutral field schema.

    Env vars required:
      MINDEE_API_KEY — Mindee API key (https://platform.mindee.com)

    API endpoint:
      POST https://api.mindee.net/v1/products/mindee/purchase_receipts/v1/predict

    On failure, returns {} so the caller can fall back to RegexOCRService.
    """

    _ENDPOINT_RECEIPTS = (
        "https://api.mindee.net/v1/products/mindee/purchase_receipts/v1/predict"
    )
    _ENDPOINT_INVOICES = (
        "https://api.mindee.net/v1/products/mindee/invoices/v4/predict"
    )

    def __init__(self) -> None:
        self._api_key = os.environ.get("MINDEE_API_KEY", "").strip()
        if not self._api_key:
            raise ValueError(
                "MINDEE_API_KEY not configured. "
                "Set OCR_PROVIDER=regex to use the default regex extractor, "
                "or set MINDEE_API_KEY in your .env file."
            )

    # ── Internal helper ───────────────────────────────────────────────────────

    def _post_to_mindee(self, endpoint: str, file_path: Optional[str], raw_text: str) -> dict:
        """
        POST document to Mindee API.
        Prefers file_path for binary upload; falls back to raw_text as a
        plain-text document if no file is provided.

        Returns the parsed JSON response dict, or {} on failure.
        """
        import requests  # requests is in requirements.txt

        headers = {"Authorization": f"Token {self._api_key}"}

        try:
            if file_path and os.path.isfile(file_path):
                with open(file_path, "rb") as fh:
                    files = {"document": fh}
                    response = requests.post(
                        endpoint, headers=headers, files=files, timeout=30
                    )
            else:
                # Send raw text as a plain .txt file
                files = {
                    "document": ("document.txt", raw_text.encode("utf-8"), "text/plain")
                }
                response = requests.post(
                    endpoint, headers=headers, files=files, timeout=30
                )

            if response.status_code == 200:
                return response.json()

            logger.warning(
                "[MindeeOCRService] API returned %s: %s",
                response.status_code,
                response.text[:300],
            )
            return {}

        except Exception as exc:
            logger.warning("[MindeeOCRService] API call failed: %s", exc)
            return {}

    # ── Field mapping helpers ─────────────────────────────────────────────────

    @staticmethod
    def _val(prediction: dict, key: str) -> Optional[str]:
        """Safely extract .value from a Mindee prediction field dict."""
        field = prediction.get(key, {})
        if isinstance(field, dict):
            v = field.get("value")
            return str(v) if v is not None else None
        if isinstance(field, list) and field:
            v = field[0].get("value") if isinstance(field[0], dict) else field[0]
            return str(v) if v is not None else None
        return None

    # ── PO extraction ─────────────────────────────────────────────────────────

    def extract_po(self, raw_text: str, file_path: Optional[str] = None) -> dict:
        """
        Extract PO fields via Mindee Receipts API.
        Maps: supplier_name → vendor, date → po_date, total_amount → total_amount,
              locale.currency → currency.
        """
        response = self._post_to_mindee(self._ENDPOINT_RECEIPTS, file_path, raw_text)
        if not response:
            return {}

        prediction = (
            response.get("document", {})
            .get("inference", {})
            .get("prediction", {})
        )
        if not prediction:
            logger.warning("[MindeeOCRService] No prediction in response for PO extraction")
            return {}

        fields: Dict[str, Any] = {}

        supplier = self._val(prediction, "supplier_name")
        if supplier:
            fields["vendor"] = supplier

        date_val = self._val(prediction, "date")
        if date_val:
            fields["po_date"] = date_val

        total = self._val(prediction, "total_amount")
        if total:
            fields["total_amount"] = total

        locale = prediction.get("locale", {})
        currency = locale.get("currency") if isinstance(locale, dict) else None
        if currency:
            fields["currency"] = str(currency).upper()

        # Line items
        line_items = prediction.get("line_items", [])
        if isinstance(line_items, list) and line_items:
            fields["line_items"] = line_items
            fields["estimated_line_count"] = len(line_items)

        logger.debug("[MindeeOCRService] extract_po — fields: %s", list(fields.keys()))
        return fields

    # ── Invoice extraction ────────────────────────────────────────────────────

    def extract_invoice(self, raw_text: str, file_path: Optional[str] = None) -> dict:
        """
        Extract invoice fields via Mindee Invoices API.
        Maps Mindee invoice prediction to neutral invoice schema.
        """
        response = self._post_to_mindee(self._ENDPOINT_INVOICES, file_path, raw_text)
        if not response:
            return {}

        prediction = (
            response.get("document", {})
            .get("inference", {})
            .get("prediction", {})
        )
        if not prediction:
            logger.warning("[MindeeOCRService] No prediction in response for invoice extraction")
            return {}

        fields: Dict[str, Any] = {}

        inv_no = self._val(prediction, "invoice_number")
        if inv_no:
            fields["invoice_number"] = inv_no

        po_ref = self._val(prediction, "reference_numbers")
        if po_ref:
            fields["po_number"] = po_ref
            fields["po_reference"] = po_ref

        supplier = self._val(prediction, "supplier_name")
        if supplier:
            fields["vendor"] = supplier

        inv_date = self._val(prediction, "date")
        if inv_date:
            fields["invoice_date"] = inv_date

        due_date = self._val(prediction, "due_date")
        if due_date:
            fields["due_date"] = due_date

        total = self._val(prediction, "total_amount")
        if total:
            fields["total_amount"] = total
            fields["amount"] = total

        tax = self._val(prediction, "total_tax")
        if tax:
            fields["tax_amount"] = tax

        locale = prediction.get("locale", {})
        currency = locale.get("currency") if isinstance(locale, dict) else None
        if currency:
            fields["currency"] = str(currency).upper()

        payment_terms = self._val(prediction, "payment_terms")
        if payment_terms:
            fields["payment_terms"] = payment_terms

        line_items = prediction.get("line_items", [])
        if isinstance(line_items, list) and line_items:
            fields["line_items"] = line_items

        logger.debug("[MindeeOCRService] extract_invoice — fields: %s", list(fields.keys()))
        return fields


# ── Implementation 3: AWSTextractService ─────────────────────────────────────

class AWSTextractService(IOCRService):
    """
    AWS Textract-based document extraction using boto3.

    Uses detect_document_text() for plain text blocks and analyze_document()
    with FORMS/TABLES features for structured field extraction.

    Env vars required:
      AWS_ACCESS_KEY_ID       — AWS access key
      AWS_SECRET_ACCESS_KEY   — AWS secret access key
      AWS_REGION              — AWS region (default: us-east-1)

    If boto3 is not installed, raises ImportError with install instructions.
    On API failure, logs warning and returns {} for graceful fallback.
    """

    def __init__(self) -> None:
        try:
            import boto3  # noqa: F401
        except ImportError:
            raise ImportError(
                "boto3 is required for AWSTextractService. "
                "Install it with: pip install boto3\n"
                "Or set OCR_PROVIDER=regex to use the default regex extractor."
            )

        self._region = os.environ.get("AWS_REGION", "us-east-1")
        self._access_key = os.environ.get("AWS_ACCESS_KEY_ID", "")
        self._secret_key = os.environ.get("AWS_SECRET_ACCESS_KEY", "")

        import boto3
        self._client = boto3.client(
            "textract",
            region_name=self._region,
            aws_access_key_id=self._access_key or None,
            aws_secret_access_key=self._secret_key or None,
        )

    # ── Internal helpers ──────────────────────────────────────────────────────

    def _get_document_bytes(self, file_path: Optional[str], raw_text: str) -> bytes:
        """Return document bytes — from file if available, else encode text."""
        if file_path and os.path.isfile(file_path):
            with open(file_path, "rb") as fh:
                return fh.read()
        return (raw_text or "").encode("utf-8")

    def _detect_text_blocks(self, doc_bytes: bytes) -> str:
        """Run detect_document_text and join all LINE blocks into a string."""
        try:
            response = self._client.detect_document_text(
                Document={"Bytes": doc_bytes}
            )
            lines = [
                block["Text"]
                for block in response.get("Blocks", [])
                if block.get("BlockType") == "LINE"
            ]
            return "\n".join(lines)
        except Exception as exc:
            logger.warning("[AWSTextractService] detect_document_text failed: %s", exc)
            return ""

    def _extract_key_value_pairs(self, doc_bytes: bytes) -> Dict[str, str]:
        """
        Run analyze_document with FORMS feature to get key-value pairs.
        Returns a flat dict: {key_text: value_text}.
        """
        try:
            response = self._client.analyze_document(
                Document={"Bytes": doc_bytes},
                FeatureTypes=["FORMS"],
            )
        except Exception as exc:
            logger.warning("[AWSTextractService] analyze_document failed: %s", exc)
            return {}

        blocks = {b["Id"]: b for b in response.get("Blocks", [])}
        kv_pairs: Dict[str, str] = {}

        for block in blocks.values():
            if block.get("BlockType") != "KEY_VALUE_SET":
                continue
            if "KEY" not in block.get("EntityTypes", []):
                continue

            # Extract key text
            key_text = _textract_collect_text(block, blocks)

            # Find value block
            for rel in block.get("Relationships", []):
                if rel["Type"] == "VALUE":
                    for val_id in rel["Ids"]:
                        val_block = blocks.get(val_id, {})
                        val_text = _textract_collect_text(val_block, blocks)
                        if key_text:
                            kv_pairs[key_text.lower().strip()] = val_text.strip()

        return kv_pairs

    # ── PO extraction ─────────────────────────────────────────────────────────

    def extract_po(self, raw_text: str, file_path: Optional[str] = None) -> dict:
        """
        Extract PO fields from AWS Textract output.
        First attempts FORMS key-value analysis, then falls back to raw text
        passed through the RegexOCRService patterns.
        """
        doc_bytes = self._get_document_bytes(file_path, raw_text)
        kv = self._extract_key_value_pairs(doc_bytes)
        full_text = self._detect_text_blocks(doc_bytes) or raw_text

        fields: Dict[str, Any] = {}

        # Map KV pairs → neutral fields
        _kv_map_po(kv, fields)

        # Fill any gaps with regex on full text
        if full_text:
            regex_fields = RegexOCRService().extract_po(full_text)
            for k, v in regex_fields.items():
                if k not in fields:
                    fields[k] = v

        logger.debug("[AWSTextractService] extract_po — fields: %s", list(fields.keys()))
        return fields

    # ── Invoice extraction ────────────────────────────────────────────────────

    def extract_invoice(self, raw_text: str, file_path: Optional[str] = None) -> dict:
        """
        Extract invoice fields from AWS Textract output.
        First attempts FORMS key-value analysis, then falls back to regex.
        """
        doc_bytes = self._get_document_bytes(file_path, raw_text)
        kv = self._extract_key_value_pairs(doc_bytes)
        full_text = self._detect_text_blocks(doc_bytes) or raw_text

        fields: Dict[str, Any] = {}

        # Map KV pairs → neutral fields
        _kv_map_invoice(kv, fields)

        # Fill any gaps with regex on full text
        if full_text:
            regex_fields = RegexOCRService().extract_invoice(full_text)
            for k, v in regex_fields.items():
                if k not in fields:
                    fields[k] = v

        logger.debug("[AWSTextractService] extract_invoice — fields: %s", list(fields.keys()))
        return fields


# ── Textract helper functions ─────────────────────────────────────────────────

def _textract_collect_text(block: dict, all_blocks: dict) -> str:
    """Collect WORD block text from a KEY or VALUE block's CHILD relationships."""
    words = []
    for rel in block.get("Relationships", []):
        if rel["Type"] == "CHILD":
            for child_id in rel["Ids"]:
                child = all_blocks.get(child_id, {})
                if child.get("BlockType") == "WORD":
                    words.append(child.get("Text", ""))
    return " ".join(words)


def _kv_map_po(kv: Dict[str, str], fields: dict) -> None:
    """Map Textract key-value pairs to PO field names."""
    _kv_copy(kv, fields, ["purchase order", "po number", "po no", "order number"], "po_number")
    _kv_copy(kv, fields, ["vendor", "supplier", "from", "vendor name"], "vendor")
    _kv_copy(kv, fields, ["date", "order date", "po date"], "po_date")
    _kv_copy(kv, fields, ["total", "total amount", "grand total", "amount"], "total_amount")
    _kv_copy(kv, fields, ["currency"], "currency")


def _kv_map_invoice(kv: Dict[str, str], fields: dict) -> None:
    """Map Textract key-value pairs to invoice field names."""
    _kv_copy(kv, fields, ["invoice number", "invoice no", "invoice #", "invoice num"], "invoice_number")
    _kv_copy(kv, fields, ["po number", "po no", "purchase order", "po reference"], "po_number")
    _kv_copy(kv, fields, ["vendor", "supplier", "billed by", "from"], "vendor")
    _kv_copy(kv, fields, ["invoice date", "date"], "invoice_date")
    _kv_copy(kv, fields, ["due date", "payment due"], "due_date")
    _kv_copy(kv, fields, ["total", "total amount", "amount due", "total due"], "total_amount")
    _kv_copy(kv, fields, ["tax", "vat", "gst"], "tax_amount")
    _kv_copy(kv, fields, ["currency"], "currency")
    _kv_copy(kv, fields, ["payment terms", "terms"], "payment_terms")

    if "total_amount" in fields:
        fields["amount"] = fields["total_amount"]
    if "po_number" in fields:
        fields["po_reference"] = fields["po_number"]


def _kv_copy(kv: Dict[str, str], fields: dict, keys: List[str], target: str) -> None:
    """Copy the first matching key from kv dict into fields[target]."""
    if target in fields:
        return
    for k in keys:
        if k in kv and kv[k]:
            fields[target] = kv[k]
            return


# ── Factory function ──────────────────────────────────────────────────────────

def get_ocr_service() -> IOCRService:
    """
    Return the configured OCR service implementation.

    Reads the OCR_PROVIDER environment variable:
      'regex'    → RegexOCRService (default, no API key needed)
      'mindee'   → MindeeOCRService (MINDEE_API_KEY required)
      'textract' → AWSTextractService (AWS credentials required)

    Falls back to RegexOCRService if the provider name is unrecognised.

    Environment Variables
    ---------------------
    OCR_PROVIDER        : provider selection (default: 'regex')
    MINDEE_API_KEY      : required for 'mindee' provider
    AWS_ACCESS_KEY_ID   : required for 'textract' provider
    AWS_SECRET_ACCESS_KEY: required for 'textract' provider
    AWS_REGION          : AWS region for 'textract' (default: 'us-east-1')
    """
    provider = os.environ.get("OCR_PROVIDER", "regex").strip().lower()

    if provider == "mindee":
        logger.info("[get_ocr_service] Using MindeeOCRService")
        return MindeeOCRService()

    if provider == "textract":
        logger.info("[get_ocr_service] Using AWSTextractService")
        return AWSTextractService()

    if provider != "regex":
        logger.warning(
            "[get_ocr_service] Unknown OCR_PROVIDER '%s'; falling back to regex.",
            provider,
        )

    logger.info("[get_ocr_service] Using RegexOCRService (default)")
    return RegexOCRService()


# ── Sprint 8: Real file-based OCR functions ───────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract plain text from a PDF document using pdfplumber.

    Returns the full concatenated text from all pages.
    Falls back to an empty string with a warning if pdfplumber is unavailable
    or the PDF cannot be read.
    """
    try:
        import pdfplumber
        import io

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages_text = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages_text.append(text)
            combined = "\n".join(pages_text)
            logger.debug(
                "[extract_text_from_pdf] Extracted %d chars from %d pages",
                len(combined), len(pdf.pages),
            )
            return combined
    except ImportError:
        logger.warning(
            "[extract_text_from_pdf] pdfplumber not installed. "
            "Run: pip install pdfplumber"
        )
        return ""
    except Exception as exc:
        logger.warning("[extract_text_from_pdf] Failed to read PDF: %s", exc)
        return ""


def extract_text_from_image(file_bytes: bytes) -> str:
    """
    Extract text from an image (PNG, JPG, TIFF, BMP) using pytesseract + PIL.

    Returns the OCR'd text string.
    Falls back to a placeholder with a warning if Tesseract binary is not
    installed or pytesseract / Pillow are unavailable.
    """
    try:
        import pytesseract
        from PIL import Image
        import io

        image = Image.open(io.BytesIO(file_bytes))
        text = pytesseract.image_to_string(image)
        logger.debug(
            "[extract_text_from_image] Extracted %d chars via OCR", len(text)
        )
        return text
    except ImportError as exc:
        logger.warning(
            "[extract_text_from_image] pytesseract / Pillow not installed (%s). "
            "Run: pip install pytesseract Pillow", exc,
        )
        return "[OCR unavailable: pytesseract/Pillow not installed]"
    except Exception as exc:
        # Tesseract binary not on PATH produces a TesseractNotFoundError
        exc_name = type(exc).__name__
        if "TesseractNotFound" in exc_name or "tesseract" in str(exc).lower():
            logger.warning(
                "[extract_text_from_image] Tesseract binary not found. "
                "Install tesseract-ocr and ensure it is on PATH. Error: %s", exc,
            )
            return "[OCR unavailable: tesseract-ocr binary not installed]"
        logger.warning("[extract_text_from_image] OCR failed: %s", exc)
        return ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract plain text from a DOCX file using python-docx.

    Returns all paragraph text joined by newlines.
    Falls back to an empty string with a warning if python-docx is unavailable.
    """
    try:
        import docx
        import io

        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        combined = "\n".join(paragraphs)
        logger.debug(
            "[extract_text_from_docx] Extracted %d chars from DOCX", len(combined)
        )
        return combined
    except ImportError:
        logger.warning(
            "[extract_text_from_docx] python-docx not installed. "
            "Run: pip install python-docx"
        )
        return ""
    except Exception as exc:
        logger.warning("[extract_text_from_docx] Failed to read DOCX: %s", exc)
        return ""


def detect_doc_type(raw_text: str) -> tuple:
    """
    Detect procurement document type from raw text using keyword scoring.

    Returns (doc_type: str, confidence: float).
    Supported types: invoice, purchase_order, delivery_note, contract, quote.
    """
    _KEYWORDS: Dict[str, List[str]] = {
        "invoice": [
            "invoice", "inv #", "inv-", "tax invoice", "amount due",
            "payment due", "remit to", "bill to", "invoice number",
        ],
        "purchase_order": [
            "purchase order", "p.o.", "po number", "po #", "ordered by",
            "order total", "ship to", "delivery address", "order date",
        ],
        "delivery_note": [
            "delivery note", "packing list", "goods received", "grn",
            "delivery order", "shipped", "consignment", "waybill",
        ],
        "contract": [
            "contract", "agreement", "terms and conditions", "obligations",
            "governing law", "indemnification", "whereas", "parties agree",
        ],
        "quote": [
            "quotation", "quote", "rfq", "request for quotation", "bid",
            "proposal", "valid until", "quoted price", "our offer",
        ],
    }

    lower = raw_text.lower()
    scores: Dict[str, int] = {}
    for doc_type, keywords in _KEYWORDS.items():
        scores[doc_type] = sum(1 for kw in keywords if kw in lower)

    if not any(scores.values()):
        return ("unknown", 0.4)

    best_type = max(scores, key=lambda t: scores[t])
    total_hits = sum(scores.values())
    confidence = round(scores[best_type] / total_hits, 2) if total_hits > 0 else 0.4
    confidence = max(min(confidence, 0.97), 0.5)
    return (best_type, confidence)


# Regex patterns for structured field extraction
import re as _re

_FIELD_PATTERNS: Dict[str, Any] = {
    "invoice_number": [
        # Must have a separator (-, #, space) after INV to avoid matching "INVOICE"
        _re.compile(r'\bINV[-#]([A-Z0-9\-/]{3,20})\b', _re.I),
        _re.compile(r'invoice\s*(?:no|number|#)[:\s#]+([A-Z0-9\-/]{3,20})', _re.I),
        _re.compile(r'(?:invoice\s+no\.?|invoice\s+number)[:\s]+([A-Z0-9\-/]{3,20})', _re.I),
    ],
    "date": [
        _re.compile(r'(?:invoice\s+date|date)[:\s]+(\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4})', _re.I),
        _re.compile(r'(\d{4}[/\-\.]\d{1,2}[/\-\.]\d{1,2})'),
        _re.compile(r'(\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4})', _re.I),
    ],
    "vendor_name": [
        _re.compile(r'(?:from|supplier|vendor|billed?\s*by)[:\s]+([A-Za-z][^\n,]{2,60})', _re.I),
    ],
    "total_amount": [
        _re.compile(r'(?:grand\s+total|total\s+amount\s+due|amount\s+due|invoice\s+total|total)[:\s]*(?:[A-Z]{3})?\s*([\d,]+\.?\d{0,2})', _re.I),
    ],
    "currency": [
        _re.compile(r'\b(AED|USD|EUR|GBP|SAR|INR|CAD|AUD|JPY|CNY)\b'),
    ],
    "po_reference": [
        _re.compile(r'(?:PO\s*ref(?:erence)?|purchase\s*order\s*(?:no|#|number)?)[:\s]+([A-Z0-9\-/]{3,20})', _re.I),
        _re.compile(r'\bPO[-/]([A-Z0-9\-]{3,20})\b', _re.I),
    ],
    "tax_amount": [
        _re.compile(r'(?:VAT|GST|tax)[:\s]*(?:[A-Z]{3})?\s*([\d,]+\.?\d{0,2})', _re.I),
    ],
    "due_date": [
        _re.compile(r'(?:due\s*date|payment\s*due)[:\s]+(\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4}|\d{4}[/\-\.]\d{1,2}[/\-\.]\d{1,2})', _re.I),
    ],
    "po_number": [
        _re.compile(r'(?:p\.?o\.?\s*(?:no|number|#)?|purchase\s*order)[:\s#]+([A-Z0-9\-/]{3,20})', _re.I),
        _re.compile(r'\bPO[-/]([A-Z0-9\-]{3,20})\b', _re.I),
    ],
    "contract_id": [
        _re.compile(r'(?:contract\s*(?:no|number|id|ref))[:\s]+([A-Z0-9\-/]{3,20})', _re.I),
    ],
    "start_date": [
        _re.compile(r'(?:start\s*date|commencement\s*date|effective\s*date)[:\s]+(\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4}|\d{4}[/\-\.]\d{1,2}[/\-\.]\d{1,2})', _re.I),
    ],
    "end_date": [
        _re.compile(r'(?:end\s*date|expiry\s*date|termination\s*date)[:\s]+(\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4}|\d{4}[/\-\.]\d{1,2}[/\-\.]\d{1,2})', _re.I),
    ],
    "contract_value": [
        _re.compile(r'(?:contract\s*value|total\s*value|agreement\s*value)[:\s]*(?:[A-Z]{3})?\s*([\d,]+\.?\d{0,2})', _re.I),
    ],
    "rfq_number": [
        _re.compile(r'\bRFQ[-#\s]?([A-Z0-9\-/]{3,20})\b', _re.I),
    ],
    "deadline": [
        _re.compile(r'(?:submission\s*deadline|response\s*by|deadline|due\s*by)[:\s]+(\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4})', _re.I),
    ],
}

# Fields relevant per document type
_DOC_TYPE_FIELDS: Dict[str, List[str]] = {
    "invoice": [
        "invoice_number", "date", "due_date", "vendor_name",
        "total_amount", "currency", "po_reference", "tax_amount",
    ],
    "purchase_order": [
        "po_number", "date", "vendor_name", "total_amount", "currency",
    ],
    "delivery_note": [
        "po_number", "date", "vendor_name",
    ],
    "contract": [
        "contract_id", "vendor_name", "start_date", "end_date",
        "contract_value", "currency",
    ],
    "quote": [
        "rfq_number", "date", "deadline", "vendor_name",
        "total_amount", "currency",
    ],
    "unknown": [
        "invoice_number", "po_number", "date", "vendor_name",
        "total_amount", "currency",
    ],
}


def extract_fields(raw_text: str, doc_type: str) -> List[dict]:
    """
    Extract structured fields from raw document text using regex patterns.

    Parameters
    ----------
    raw_text : str  — the full document text
    doc_type : str  — one of: invoice, purchase_order, delivery_note, contract, quote, unknown

    Returns
    -------
    list of {name, value, confidence, status} dicts.
    status is 'extracted' when value found, 'missing' otherwise.
    """
    fields_to_extract = _DOC_TYPE_FIELDS.get(doc_type, _DOC_TYPE_FIELDS["unknown"])
    results = []

    for field_name in fields_to_extract:
        patterns = _FIELD_PATTERNS.get(field_name, [])
        value = None
        confidence = 0.0

        for i, pattern in enumerate(patterns):
            m = pattern.search(raw_text)
            if m:
                value = m.group(1).strip()
                # First pattern = strongest match
                confidence = 1.0 if i == 0 else (0.85 if i == 1 else 0.7)
                break

        # Heuristic fallback for vendor name if no label found
        if field_name == "vendor_name" and not value:
            lines = [l.strip() for l in raw_text.splitlines() if l.strip()]
            if lines:
                # Take first non-trivial line as a heuristic
                for line in lines[:10]:
                    if len(line) > 5 and not any(
                        kw in line.lower() for kw in
                        ("invoice", "tax", "date", "total", "page", "bill", "from", "to:")
                    ):
                        value = line[:60]
                        confidence = 0.5
                        break

        if value:
            results.append({
                "name": field_name,
                "value": value,
                "confidence": confidence,
                "status": "extracted",
            })
        else:
            results.append({
                "name": field_name,
                "value": None,
                "confidence": 0.0,
                "status": "missing",
            })

    return results


def process_document(
    file_bytes: bytes,
    filename: str,
    doc_type: str = "auto",
) -> dict:
    """
    Full document processing pipeline: extract text → detect type → extract fields.

    Parameters
    ----------
    file_bytes : bytes  — raw file content
    filename   : str    — original filename (used to determine file type)
    doc_type   : str    — "auto" to detect, or explicit type hint

    Returns
    -------
    {
      raw_text          : str,
      doc_type_detected : str,
      confidence        : float,
      fields            : list of {name, value, confidence, status},
      source            : str,   # 'pdf', 'image', 'docx', 'text'
    }
    """
    ext = (filename.rsplit(".", 1)[-1].lower() if "." in filename else "").lower()

    # ── Step 1: Extract raw text ──────────────────────────────────────────────
    raw_text = ""
    source = "text"

    if ext == "pdf":
        raw_text = extract_text_from_pdf(file_bytes)
        source = "pdf"
    elif ext in ("png", "jpg", "jpeg", "tiff", "tif", "bmp"):
        raw_text = extract_text_from_image(file_bytes)
        source = "image"
    elif ext in ("docx", "doc"):
        raw_text = extract_text_from_docx(file_bytes)
        source = "docx"
    else:
        # Try to decode as plain text
        try:
            raw_text = file_bytes.decode("utf-8", errors="replace")
            source = "text"
        except Exception:
            raw_text = ""
            source = "unknown"

    if not raw_text.strip():
        logger.warning(
            "[process_document] No text extracted from '%s' (ext=%s)", filename, ext
        )

    # ── Step 2: Detect document type ──────────────────────────────────────────
    if doc_type == "auto" or not doc_type:
        doc_type_detected, confidence = detect_doc_type(raw_text)
    else:
        doc_type_detected = doc_type.lower()
        confidence = 1.0  # caller-supplied type is authoritative

    # ── Step 3: Extract structured fields ────────────────────────────────────
    fields = extract_fields(raw_text, doc_type_detected)

    logger.info(
        "[process_document] file='%s' source=%s type=%s confidence=%.2f fields=%d",
        filename, source, doc_type_detected, confidence, len(fields),
    )

    return {
        "raw_text": raw_text,
        "doc_type_detected": doc_type_detected,
        "confidence": confidence,
        "fields": fields,
        "source": source,
        "filename": filename,
    }
